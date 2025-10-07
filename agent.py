import os, datetime as dt, pytz
import click
from ruamel.yaml import YAML
from jinja2 import Environment, FileSystemLoader, select_autoescape
from docx import Document
from ics import Calendar, Event

yaml = YAML()

def load(yaml_path):
    with open(yaml_path, "r", encoding="utf-8") as f:
        return yaml.load(f)

def jenv():
    env = Environment(
        loader=FileSystemLoader("templates"),
        autoescape=select_autoescape()
    )
    env.globals["fmt_money"] = fmt_money
    return env

def write_docx(text, outpath):
    doc = Document()
    for para in text.split("\n"):
        doc.add_paragraph(para)
    os.makedirs(os.path.dirname(outpath), exist_ok=True)
    doc.save(outpath)

def fmt_money(x): return f"{x:,.2f}"

def compute_amount(balance, pct):
    return round(balance * pct / 100.0, 2)

@click.group()
def cli():
    """Debt/Credit Repair Agent"""

@cli.command()
@click.option("--cfg", default="data/plan.yaml")
def make_calendar(cfg):
    """Emit .ics with your deadlines."""
    plan = load(cfg)
    tz = pytz.timezone(plan["tz"])
    cal = Calendar()
    def add(uid, date_str, title):
        d = tz.localize(dt.datetime.strptime(date_str, "%Y-%m-%d"))
        ev = Event(name=title, begin=d, duration=dt.timedelta(hours=1))
        ev.uid = uid
        cal.events.add(ev)
    dd = plan["deadlines"]
    add("opp_msj", dd["oppose_msj"], "File Opposition to MSJ + Motion to Strike")
    add("mov_623", dd["mov_bureaus_623"], "Mail §611 MOV to EQ/TU/EX + §623 to Chase")
    add("fdcpas", dd["fdcpas"], "Mail FDCPA validations to SCS and Credit Collections")
    add("rebuild", dd["rebuild_day45"], "Day 45: Rebuild wave kickoff")
    out = os.path.join(plan["out_dir"], "deadlines.ics")
    os.makedirs(plan["out_dir"], exist_ok=True)
    with open(out, "w", encoding="utf-8") as f: f.writelines(cal.serialize_iter())
    click.echo(out)

@cli.command()
@click.option("--cfg", default="data/plan.yaml")
def filings(cfg):
    """Generate court filings for Chase 4155."""
    plan = load(cfg); env = jenv()
    now = dt.datetime.now().strftime("%Y-%m-%d")
    ctx = {**plan, "now": now}

    # Opposition
    text = env.get_template("opposition_msj.j2").render(**ctx)
    fn = f'{plan["court"]["cause_no"]}_Opposition_to_MSJ.docx'
    write_docx(text, os.path.join(plan["out_dir"], fn))

    # Motion to Strike
    text = env.get_template("motion_strike_affidavit.j2").render(**ctx)
    fn = f'{plan["court"]["cause_no"]}_Motion_to_Strike_Affidavit.docx'
    write_docx(text, os.path.join(plan["out_dir"], fn))

    # Hardship declaration
    text = env.get_template("hardship_declaration.j2").render(**ctx)
    fn = "Hardship_Declaration.docx"
    write_docx(text, os.path.join(plan["out_dir"], fn))

    click.echo("filings OK")

@cli.command()
@click.option("--cfg", default="data/plan.yaml")
def letters(cfg):
    """Generate all dispute and settlement letters."""
    plan = load(cfg); env = jenv()
    today = dt.datetime.now().strftime("%Y-%m-%d")
    party = plan["party"]; facts = plan["facts"]

    out_dir = plan["out_dir"]; os.makedirs(out_dir, exist_ok=True)

    # Settlement – Chase 4155
    for i, pct in enumerate(plan["offers"]["chase_4155"], start=1):
        ctx = {
            **plan, "date": today, "name": "Chase 4155",
            "balance": fmt_money(facts["chase_balance_4155"]),
            "pct": pct, "amount": fmt_money(compute_amount(facts["chase_balance_4155"], pct)),
            "to_address": plan["mailing"]["plaintiff_counsel"], "wave": i
        }
        text = env.get_template("settlement_offer.j2").render(**ctx)
        fn = f"Chase_4155_Wave{i}_Settlement_Offer_{today}.docx"
        write_docx(text, os.path.join(out_dir, fn))

    # §611 MOV to bureaus (duplicate JPMCB *4978)
    for bureau in plan["mailing"]["bureaus"]:
        ctx = {**plan, "date": today, "to_address": bureau}
        text = env.get_template("frca_611_mov.j2").render(**ctx)
        fn = f'Bureaus_§611_MOV_JPMCB4978_{today}_{bureau.split(",")[0].replace(" ", "")}.docx'
        write_docx(text, os.path.join(out_dir, fn))

    # §623 direct dispute to Chase
    text = env.get_template("frca_623_direct.j2").render(
        **plan, date=today, to_address=plan["mailing"]["chase_furnisher"]
    )
    fn = f"Chase_§623_DirectDispute_DuplicateTradeline_{today}.docx"
    write_docx(text, os.path.join(out_dir, fn))

    # FDCPA validations + PFD ladders
    # SCS (ComEd)
    for i, pct in enumerate(plan["offers"]["scs_comed"], start=1):
        ctx = {
            **plan, "date": today, "pct": pct,
            "amount": fmt_money(compute_amount(facts["scs_comed_balance"], pct)),
            "to_address": plan["mailing"]["scs"][0],
            "account_reference": f"Southwest Credit Systems – ComEd Balance ${fmt_money(facts['scs_comed_balance'])}"
        }
        text = env.get_template("fdCPA_1692g_pfd.j2").render(**ctx)
        fn = f"SCS_ComEd_Validation_PFD_{today}_Wave{i}.docx"
        write_docx(text, os.path.join(out_dir, fn))

    # Credit Collections (Progressive)
    for i, pct in enumerate(plan["offers"]["credit_collections_prog"], start=1):
        ctx = {
            **plan, "date": today, "pct": pct,
            "amount": fmt_money(compute_amount(facts["credit_collections_prog_balance"], pct)),
            "to_address": plan["mailing"]["credit_collections"][0],
            "account_reference": f"Credit Collection Services – Progressive Balance ${fmt_money(facts['credit_collections_prog_balance'])}"
        }
        text = env.get_template("fdCPA_1692g_pfd.j2").render(**ctx)
        fn = f"CreditCollections_Progressive_Validation_PFD_{today}_Wave{i}.docx"
        write_docx(text, os.path.join(out_dir, fn))

    # AmEx + BofA offers
    for name, bal_key, ladder in [
        ("Amex_15443", "amex_balance", plan["offers"]["amex"]),
        ("BofA", "bofa_balance", plan["offers"]["bofa"])
    ]:
        for i, pct in enumerate(ladder, start=1):
            ctx = {
                **plan, "date": today, "name": name,
                "balance": fmt_money(facts[bal_key]),
                "pct": pct, "amount": fmt_money(compute_amount(facts[bal_key], pct)),
                "to_address": "[Creditor address here]", "wave": i
            }
            text = env.get_template("settlement_offer.j2").render(**ctx)
            fn = f"{name}_Wave{i}_Settlement_Offer_{today}.docx"
            write_docx(text, os.path.join(out_dir, fn))

    click.echo("letters OK")

@cli.command()
@click.option("--cfg", default="data/plan.yaml")
def scaffold(cfg):
    """Create out/ filenames exactly as specified in your brief."""
    plan = load(cfg)
    out = plan["out_dir"]; os.makedirs(out, exist_ok=True)
    court = plan["court"]["cause_no"]
    today = dt.datetime.now().strftime("%Y-%m-%d")
    targets = [
        f"{court}_Opposition_to_MSJ.docx",
        f"{court}_Motion_to_Strike_Affidavit.docx",
        f"Chase_4155_Wave1_Settlement_Offer_{today}.docx",
        f"Bureaus_§611_MOV_JPMCB4978_{today}.docx",
        f"Chase_§623_DirectDispute_DuplicateTradeline_{today}.docx",
        f"SCS_ComEd_Validation_PFD_{today}.docx",
        f"CreditCollections_Validation_PFD_Progressive_{today}.docx",
        f"Amex_15443_Wave1_Settlement_Offer_{today}.docx",
        f"BofA_Itemization_and_Settlement_Wave1_{today}.docx",
    ]
    for t in targets:
        p = os.path.join(out, t)
        if not os.path.exists(p):
            Document().save(p)
    click.echo("scaffold OK")

if __name__ == "__main__":
    cli()
